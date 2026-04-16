
# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

OUTPUT_PATH = r'التوثيق الاخير 1 - منسق.docx'
doc = Document(OUTPUT_PATH)

print("=" * 60)
print("VERIFICATION REPORT")
print("=" * 60)

tables = doc.tables
print(f"\nTotal tables: {len(tables)}\n")

# ── Identify key tables by heading scan ──
all_elements = list(doc.element.body)
para_indices = {}
tbl_indices  = {}
for idx, el in enumerate(all_elements):
    tag = el.tag.split('}')[-1]
    if tag == 'p':   para_indices[id(el)] = idx
    elif tag == 'tbl': tbl_indices[id(el)] = idx

ordered = []
for p in doc.paragraphs:
    ordered.append((para_indices.get(id(p._element), 0), 'para', p))
for t in doc.tables:
    ordered.append((tbl_indices.get(id(t._tbl), 0), 'table', t))
ordered.sort(key=lambda x: x[0])

last_heading = ''
table_summary = {}
for idx, typ, obj in ordered:
    if typ == 'para':
        txt = obj.text.strip()
        if txt: last_heading = txt
    elif typ == 'table':
        h = last_heading
        rows = len(obj.rows)
        cols = len(obj.columns)
        table_summary[h[:40]] = (rows, cols, obj)

print("KEY TABLES FOUND:")
print("-" * 60)
for heading, (rows, cols, t) in table_summary.items():
    print(f"  [{heading}] -> rows={rows}, cols={cols}")

# ── Check TOC ──
print("\n─── TABLE OF CONTENTS (first 10 entries) ─────────────")
for heading, (rows, cols, t) in table_summary.items():
    if 'محتوي' in heading or 'محتويات' in heading:
        for i, row in enumerate(t.rows[:11]):
            cells = [c.text.strip()[:50] for c in row.cells]
            print(f"  Row {i}: {' | '.join(cells)}")
        print(f"  ... total {rows} rows")
        break

# ── Check Figures ──
print("\n─── FIGURES INDEX (first 5 entries) ──────────────────")
for heading, (rows, cols, t) in table_summary.items():
    if 'أشكال' in heading or 'اشكال' in heading:
        for i, row in enumerate(t.rows[:6]):
            cells = [c.text.strip()[:50] for c in row.cells]
            print(f"  Row {i}: {' | '.join(cells)}")
        print(f"  ... total {rows} rows")
        break

# ── Check Tables Index ──
print("\n─── TABLES INDEX (first 5 entries) ───────────────────")
for heading, (rows, cols, t) in table_summary.items():
    if 'جداول' in heading and 'أشكال' not in heading and 'اشكال' not in heading:
        for i, row in enumerate(t.rows[:6]):
            cells = [c.text.strip()[:50] for c in row.cells]
            print(f"  Row {i}: {' | '.join(cells)}")
        print(f"  ... total {rows} rows")
        break

# ── Check Abbreviations ──
print("\n─── ABBREVIATIONS (all rows) ──────────────────────────")
abbr_t = None
for heading, (rows, cols, t) in table_summary.items():
    hdr = t.rows[0].cells[0].text.strip() if t.rows else ''
    if 'اختصار' in hdr or 'Abbreviation' in hdr.lower():
        abbr_t = t
        break
if abbr_t is None:
    abbr_t = doc.tables[-1]

for i, row in enumerate(abbr_t.rows):
    cells = [c.text.strip()[:30] for c in row.cells]
    print(f"  Row {i}: {' | '.join(cells)}")
print(f"  Total: {len(abbr_t.rows)} rows")

# ── Paragraph formatting sample ──
print("\n─── PARAGRAPH FORMATTING SAMPLE ──────────────────────")
count = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt or count > 20:
        continue
    # Only print headings and first body lines
    size = None
    bold = None
    if p.runs:
        s = p.runs[0].font.size
        b = p.runs[0].font.bold
        if s: size = round(s/12700, 1)
        bold = b
    if size and size >= 14:
        align_name = str(p.alignment).replace('WD_ALIGN_PARAGRAPH.', '') if p.alignment else 'None'
        print(f"  [{size}pt bold={bold} align={align_name}] {txt[:60]}")
        count += 1

print("\n" + "=" * 60)
print("Verification complete.")
